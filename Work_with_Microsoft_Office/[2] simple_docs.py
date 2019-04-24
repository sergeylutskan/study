from docx import Document

def markdown_to_docx(text):
    document = Document()
    lines = text.split('\n')
    document.add_heading(lines[0],0)
    for line in lines[1:]:
        if line:
            if line[:7].count('#') == 1:
                document.add_heading(line[1:].replace('#',''), level=1)
            elif line[:7].count('#') == 2:
                document.add_heading(line[2:].replace('#',''), level=2)
            elif line[:7].count('#') == 3:
                document.add_heading(line[3:].replace('#',''), level=3)
            elif line[:7].count('#') == 4:
                document.add_heading(line[4:].replace('#',''), level=4)
            elif line[:7].count('#') == 5:
                document.add_heading(line[5:].replace('#',''), level=5)
            elif line[:7].count('#') == 6:
                document.add_heading(line[6:].replace('#',''), level=6)
            elif str(line[:2]) == '- ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '* ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif str(line[:2]) == '+ ':
                document.add_paragraph(line[2:], style='List Bullet')
            elif line[0].isdigit() and line[1] == '.':
                document.add_paragraph(line[2:], style='List Number')
            elif line[:2].count('*') == 1:
                document.add_paragraph().add_run(line[1:]).italic = True
            elif line[:2].count('*') == 2:
                document.add_paragraph().add_run(line[2:]).bold = True
            else:
                document.add_paragraph(line)
        else:
            document.add_paragraph()
    document.save('test.docx')

markdown_to_docx('''test01
Абзацы создаются при помощи пустой строки. Если вокруг текста сверху и снизу есть пустые строки, то текст превращается в абзац.
*Чтобы сделать перенос строки вместо абзаца,
нужно поставить два пробела в конце предыдущей строки.
**Заголовки отмечаются диезом `#` в начале строки, от одного до шести. Например:
# Заголовок первого уровня #
## Заголовок h2
### Заголовок h3
#### Заголовок h4
##### Заголовок h5
###### Заголовок h6
В декоративных целях заголовки можно «закрывать» с обратной стороны.
### Списки
Для разметки неупорядоченных списков можно использовать или `*`, или `-`, или `+`:
- элемент 1
- элемент 2
- элемент ...

Упорядоченный список:
1. элемент 1
2. элемент 2
3. элемент 3
4. Donec sit amet nisl. Aliquam semper ipsum sit amet velit. Suspendisse id sem consectetuer libero luctus adipiscing.
На самом деле не важно как в коде пронумерованы пункты, главное, чтобы перед элементом списка стояла цифра (любая) с точкой. Можно сделать и так:
0. элемент 1
0. элемент 2
0. элемент 3
0. элемент 4

''')
